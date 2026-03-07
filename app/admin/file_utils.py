"""Utilities for extracting plain text from uploaded files (PDF, DOCX, TXT)."""

from docx import Document as DocxDocument
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage


def extract_text_from_pdf(file_storage: FileStorage) -> str:
    """Extract text from all pages of a PDF file."""
    reader = PdfReader(file_storage.stream)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def extract_text_from_docx(file_storage: FileStorage) -> str:
    """Extract text from paragraphs and tables of a DOCX file."""
    doc = DocxDocument(file_storage.stream)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_file(file_storage: FileStorage, filename: str) -> str:
    """Dispatch to the correct extractor based on file extension.

    Raises ValueError for unsupported types or empty content.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        text = extract_text_from_pdf(file_storage)
    elif ext == "docx":
        text = extract_text_from_docx(file_storage)
    elif ext == "txt":
        text = file_storage.stream.read().decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    text = text.strip()
    if not text:
        raise ValueError("File appears to be empty — no text could be extracted.")

    return text
