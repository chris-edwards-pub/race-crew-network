"""Tests for file text extraction utilities."""

import io

import pytest
from werkzeug.datastructures import FileStorage

from app.admin.file_utils import extract_text_from_file


class TestExtractTextFromFile:
    """Tests for the extract_text_from_file dispatcher."""

    def test_txt_passthrough(self):
        content = b"Regatta schedule\nEvent 1\nEvent 2"
        fs = FileStorage(stream=io.BytesIO(content), filename="schedule.txt")
        result = extract_text_from_file(fs, "schedule.txt")
        assert "Regatta schedule" in result
        assert "Event 1" in result

    def test_unsupported_type_raises(self):
        fs = FileStorage(stream=io.BytesIO(b"data"), filename="image.jpg")
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text_from_file(fs, "image.jpg")

    def test_empty_txt_raises(self):
        fs = FileStorage(stream=io.BytesIO(b""), filename="empty.txt")
        with pytest.raises(ValueError, match="empty"):
            extract_text_from_file(fs, "empty.txt")

    def test_whitespace_only_txt_raises(self):
        fs = FileStorage(stream=io.BytesIO(b"   \n\n  "), filename="blank.txt")
        with pytest.raises(ValueError, match="empty"):
            extract_text_from_file(fs, "blank.txt")

    def test_pdf_empty_raises(self):
        """A blank-page PDF with no extractable text raises ValueError."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)

        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)

        fs = FileStorage(stream=buf, filename="blank.pdf")
        with pytest.raises(ValueError, match="empty"):
            extract_text_from_file(fs, "blank.pdf")

    def test_pdf_extraction_with_text(self):
        """Test PDF extraction using mock to verify the pipeline."""
        from unittest.mock import patch

        fake_content = "Regatta Schedule 2026"
        fs = FileStorage(stream=io.BytesIO(b"fake"), filename="schedule.pdf")

        with patch(
            "app.admin.file_utils.extract_text_from_pdf", return_value=fake_content
        ):
            result = extract_text_from_file(fs, "schedule.pdf")
        assert result == fake_content

    def test_docx_extraction(self):
        """Test DOCX extraction with a minimal valid DOCX."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Spring Regatta Schedule")
        doc.add_paragraph("Event 1: Midwinters - March 15")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        fs = FileStorage(stream=buf, filename="schedule.docx")
        result = extract_text_from_file(fs, "schedule.docx")
        assert "Spring Regatta Schedule" in result
        assert "Midwinters" in result

    def test_docx_with_tables(self):
        """Test DOCX extraction includes table content."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Schedule")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Event"
        table.cell(0, 1).text = "Date"
        table.cell(1, 0).text = "Midwinters"
        table.cell(1, 1).text = "March 15"

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        fs = FileStorage(stream=buf, filename="schedule.docx")
        result = extract_text_from_file(fs, "schedule.docx")
        assert "Schedule" in result
        assert "Midwinters" in result
        assert "March 15" in result

    def test_xlsx_extraction(self):
        """Test Excel extraction with a minimal valid XLSX."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Event", "Date", "Location"])
        ws.append(["Midwinters", "March 15", "Eustis Sailing Club"])
        ws.append(["Spring Regatta", "April 10", "Lake Lanier SC"])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        fs = FileStorage(stream=buf, filename="schedule.xlsx")
        result = extract_text_from_file(fs, "schedule.xlsx")
        assert "Midwinters" in result
        assert "Eustis Sailing Club" in result
        assert "Spring Regatta" in result

    def test_xlsx_skips_empty_rows(self):
        """Test Excel extraction skips rows with no data."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Event 1"])
        ws.append([None, None])
        ws.append(["Event 2"])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        fs = FileStorage(stream=buf, filename="schedule.xlsx")
        result = extract_text_from_file(fs, "schedule.xlsx")
        assert "Event 1" in result
        assert "Event 2" in result
        lines = [ln for ln in result.split("\n") if ln.strip()]
        assert len(lines) == 2

    def test_xlsx_empty_raises(self):
        """An empty Excel file raises ValueError."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append([None])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        fs = FileStorage(stream=buf, filename="empty.xlsx")
        with pytest.raises(ValueError, match="empty"):
            extract_text_from_file(fs, "empty.xlsx")

    def test_no_extension_raises(self):
        fs = FileStorage(stream=io.BytesIO(b"data"), filename="noext")
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text_from_file(fs, "noext")
